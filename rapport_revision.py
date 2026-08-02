"""
Rapport de revision au format Word, destine au client.

Le document ne reprend QUE les points appelant une action du client : anomalies
non traitees, hors controles techniques internes. Il sert de demande de
justificatifs, pas de rapport d'audit.

Sont exclus :
  - la famille MAP, qui porte sur le parametrage de l'application et ne
    concerne en rien le client
  - les anomalies deja justifiees, corrigees ou ignorees
  - les lignes informatives, sauf celles qui detaillent une anomalie retenue
"""
import io
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

BLEU = RGBColor(0x1A, 0x52, 0x76)
GRIS = RGBColor(0x56, 0x65, 0x73)

SEVERITES = ("bloquant", "majeur", "justifier")
SEV_LABEL = {
    "bloquant": "Bloquant",
    "majeur": "Majeur",
    "justifier": "A justifier",
    "informatif": "Informatif",
}
FAMILLES_EXCLUES = ("MAP",)

FAMILLE_TITRE = {
    "INT": "Integrite de la balance",
    "SEN": "Comptes au sens inhabituel",
    "ATT": "Comptes d'attente et de regularisation",
    "FIS": "Coherence fiscale",
    "SOC": "Charges sociales et paie",
    "IMM": "Immobilisations et amortissements",
    "STK": "Stocks",
    "VAR": "Variations par rapport a l'exercice precedent",
    "ETA": "Coherence des etats financiers",
    "CPL": "Informations manquantes",
}
FAMILLE_ORDRE = ["INT", "ETA", "SEN", "ATT", "IMM", "STK", "FIS", "SOC", "VAR", "CPL"]


def _fmt(x):
    if x is None:
        return ""
    try:
        return "{:,.0f}".format(float(x)).replace(",", " ")
    except (TypeError, ValueError):
        return ""


def _date_fr(valeur):
    if not valeur:
        return ""
    try:
        return datetime.strptime(str(valeur)[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return str(valeur)[:10]


def _ombrer(cellule, couleur):
    """Fond de cellule. ShadingType SOLID rend du noir dans certains lecteurs :
    on passe par un element w:shd explicite avec un remplissage clair."""
    tc = cellule._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), couleur)
    tc.append(shd)


def _bordure_basse(paragraphe):
    p = paragraphe._p.get_or_add_pPr()
    bordures = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), "8")
    bas.set(qn("w:color"), "1A5276")
    bordures.append(bas)
    p.append(bordures)


def _repeter_entete(ligne):
    """Repete la ligne d'en-tete en haut de chaque page : sans cela, un tableau
    qui se poursuit sur plusieurs pages devient illisible."""
    trPr = ligne._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _selection(anomalies):
    """Anomalies retenues pour le client, plus leurs lignes de detail."""
    retenues = [
        a for a in anomalies
        if a["famille"] not in FAMILLES_EXCLUES
        and a["severite"] in SEVERITES
        and a["statut"] == "traiter"
    ]
    codes = {a["code"] for a in retenues}
    # Les lignes de detail (suffixe D) accompagnent une anomalie principale :
    # on ne les garde que si celle-ci figure au rapport.
    details = [
        a for a in anomalies
        if a["code"].endswith("D") and a["code"][:-1] in codes
    ]
    return retenues, details


def generer(anomalies, client, exo, cabinet=None):
    """Retourne le document Word en octets."""
    retenues, details = _selection(anomalies)
    par_code = {}
    for d in details:
        par_code.setdefault(d["code"][:-1], []).append(d)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for marge in ("left_margin", "right_margin"):
        setattr(section, marge, Cm(2.0))
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── En-tete ────────────────────────────────────────────────────────────
    if cabinet:
        p = doc.add_paragraph()
        r = p.add_run(cabinet)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRIS

    titre = doc.add_paragraph()
    r = titre.add_run("RAPPORT DE REVISION COMPTABLE")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLEU
    _bordure_basse(titre)

    st = doc.add_paragraph()
    r = st.add_run("Points appelant votre attention")
    r.font.size = Pt(11)
    r.font.color.rgb = GRIS
    r.italic = True

    doc.add_paragraph()

    # ── Identification ─────────────────────────────────────────────────────
    ident = doc.add_table(rows=0, cols=2)
    ident.columns[0].width = Cm(4.5)
    ident.columns[1].width = Cm(12.5)
    lignes = [
        ("Entite", (client["raison_sociale"] if client else "") or ""),
        ("N° de compte contribuable", (client["ncc"] if client else "") or "—"),
        ("Exercice", str(exo["annee"]) if exo else ""),
    ]
    if exo and (exo["date_debut"] or exo["date_fin"]):
        lignes.append(("Periode", "%s au %s" % (_date_fr(exo["date_debut"]),
                                                _date_fr(exo["date_fin"]))))
    lignes.append(("Date d'edition", datetime.now().strftime("%d/%m/%Y")))
    for libelle, valeur in lignes:
        row = ident.add_row()
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(12.5)
        c0 = row.cells[0].paragraphs[0].add_run(libelle)
        c0.bold = True
        c0.font.size = Pt(9)
        c1 = row.cells[1].paragraphs[0].add_run(str(valeur))
        c1.font.size = Pt(9)

    doc.add_paragraph()

    # ── Introduction ───────────────────────────────────────────────────────
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if retenues:
        texte = (
            "La revision de votre balance comptable a fait apparaitre %d point%s "
            "necessitant une correction de votre part ou une explication. Vous "
            "trouverez ci-apres le detail de ces points, classes par nature. "
            "Merci de nous retourner vos justificatifs ou vos corrections afin que "
            "nous puissions finaliser vos etats financiers."
            % (len(retenues), "s" if len(retenues) > 1 else "")
        )
    else:
        texte = (
            "La revision de votre balance comptable n'a fait apparaitre aucun point "
            "necessitant une correction ou une explication de votre part. Vos etats "
            "financiers peuvent etre finalises en l'etat."
        )
    intro.add_run(texte).font.size = Pt(10)

    # ── Synthese ───────────────────────────────────────────────────────────
    if retenues:
        doc.add_paragraph()
        h = doc.add_paragraph()
        r = h.add_run("Synthese")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BLEU

        compte = {s: 0 for s in SEVERITES}
        for a in retenues:
            compte[a["severite"]] = compte.get(a["severite"], 0) + 1

        syn = doc.add_table(rows=1, cols=3)
        syn.alignment = WD_TABLE_ALIGNMENT.LEFT
        entetes = ("Niveau", "Nombre", "Signification")
        largeurs = (Cm(3.5), Cm(2.0), Cm(11.5))
        for i, (txt, larg) in enumerate(zip(entetes, largeurs)):
            cel = syn.rows[0].cells[i]
            cel.width = larg
            _ombrer(cel, "1A5276")
            run = cel.paragraphs[0].add_run(txt)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        explications = {
            "bloquant": "A corriger imperativement avant depot des etats",
            "majeur": "Anomalie probable, correction attendue",
            "justifier": "Peut etre normal, une explication est demandee",
        }
        for s in SEVERITES:
            if not compte.get(s):
                continue
            row = syn.add_row()
            for i, (txt, larg) in enumerate(zip(
                    (SEV_LABEL[s], str(compte[s]), explications[s]), largeurs)):
                row.cells[i].width = larg
                run = row.cells[i].paragraphs[0].add_run(txt)
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

    # ── Detail par famille ─────────────────────────────────────────────────
    familles = {}
    for a in retenues:
        familles.setdefault(a["famille"], []).append(a)

    ordre = [f for f in FAMILLE_ORDRE if f in familles]
    ordre += [f for f in familles if f not in FAMILLE_ORDRE]

    poids = {s: i for i, s in enumerate(SEVERITES)}
    for fam in ordre:
        lot = sorted(familles[fam], key=lambda a: (poids.get(a["severite"], 9),
                                                   a["compte"] or ""))
        doc.add_paragraph()
        h = doc.add_paragraph()
        r = h.add_run(FAMILLE_TITRE.get(fam, fam))
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = BLEU

        t = doc.add_table(rows=1, cols=4)
        largeurs = (Cm(2.0), Cm(1.8), Cm(10.2), Cm(3.0))
        for i, txt in enumerate(("Compte", "Niveau", "Constat", "Montant")):
            cel = t.rows[0].cells[i]
            cel.width = largeurs[i]
            _ombrer(cel, "EAF2F8")
            run = cel.paragraphs[0].add_run(txt)
            run.bold = True
            run.font.size = Pt(9)
            if i == 3:
                cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _repeter_entete(t.rows[0])

        for a in lot:
            row = t.add_row()
            valeurs = (a["compte"] or "—", SEV_LABEL.get(a["severite"], ""),
                       a["libelle"], _fmt(a["montant"]))
            for i, txt in enumerate(valeurs):
                cel = row.cells[i]
                cel.width = largeurs[i]
                run = cel.paragraphs[0].add_run(str(txt))
                run.font.size = Pt(8.5)
                if i == 3:
                    cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if i == 1 and a["severite"] == "bloquant":
                    run.bold = True

            # Lignes de detail rattachees a cette anomalie
            for d in par_code.get(a["code"], []):
                sr = t.add_row()
                for i, txt in enumerate(("", "", "    " + d["libelle"], _fmt(d["montant"]))):
                    cel = sr.cells[i]
                    cel.width = largeurs[i]
                    run = cel.paragraphs[0].add_run(str(txt))
                    run.font.size = Pt(8)
                    run.font.color.rgb = GRIS
                    if i == 3:
                        cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                sr.cells[0].paragraphs[0].add_run(d["compte"] or "").font.size = Pt(8)

    # ── Pied ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fin = doc.add_paragraph()
    fin.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = fin.add_run(
        "Ce rapport est etabli a partir de la balance comptable qui nous a ete "
        "transmise. Il ne constitue ni un audit ni une certification des comptes. "
        "Les montants sont exprimes en francs CFA."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GRIS
    r.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def nom_fichier(client, exo):
    raison = ((client["raison_sociale"] if client else "") or "client")
    raison = "".join(ch if ch.isalnum() else "_" for ch in raison)[:30].strip("_")
    return "Rapport_revision_%s_%s.docx" % (raison, exo["annee"] if exo else "")
